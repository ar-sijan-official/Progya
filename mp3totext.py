import sys
import os

# --- PYTHON 3.14 COMPATIBILITY FIX ---
try:
    import audioop
except ImportError:
    import audioop_lts as audioop
    sys.modules["audioop"] = audioop

import whisper
from docx import Document

def convert_audio_to_text(mp3_path, output_format="both"):
    """
    Transcribes a mixed Bangla-English MP3 file and saves it as .txt and .docx
    """
    if not os.path.exists(mp3_path):
        print(f"[!] Error: File not found at {mp3_path}")
        return

    print(f"[*] Loading Whisper 'small' model (Better for Bangla/English mix)...")
    # Using 'small' because 'base' often struggles with Bangla script accuracy
    model = whisper.load_model("small")

    print(f"[*] Transcribing: {os.path.basename(mp3_path)}...")
    
    # We use an 'initial_prompt' to tell the AI to expect both languages.
    # This helps keep technical English terms in English script.
    benglish_prompt = "This is a classroom lecture. The teacher speaks in Bangla and English technical terms like Python, loop, and functions."

    result = model.transcribe(
        mp3_path, 
        initial_prompt=benglish_prompt,
        verbose=False
    )
    
    text = result["text"]

    # Define base filename (removes .mp3 extension)
    base_name = os.path.splitext(mp3_path)[0]

    # Save as .txt
    if output_format in ["txt", "both"]:
        txt_path = f"{base_name}.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"[+] Saved Text: {txt_path}")

    # Save as .docx
    if output_format in ["docx", "both"]:
        docx_path = f"{base_name}.docx"
        doc = Document()
        doc.add_heading('Classroom Lecture Transcription', 0)
        doc.add_paragraph(text)
        doc.save(docx_path)
        print(f"[+] Saved Word Doc: {docx_path}")

if __name__ == "__main__":
    # Your specific file location
    target_file = "/home/zatch/Desktop/Sentinel/ai_class/mp3_recordings/lecture_2026-02-08_01-29-17.mp3"
    
    convert_audio_to_text(target_file)